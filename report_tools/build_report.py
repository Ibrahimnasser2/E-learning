"""Generate the Arabic Word report (MANAMU presentation) using python-docx."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "screenshots"
OUT = ROOT.parent / "MANAMU_Presentation_Report.docx"

# (file_name, title, description bullets)
SECTIONS = [
    (
        "01_login.png",
        "1. صفحة تسجيل الدخول (Login)",
        [
            "نقطة الدخول الرئيسية للنظام لكل من الطلاب وأعضاء هيئة التدريس.",
            "الحقول المطلوبة: اسم المستخدم، كلمة المرور، اختيار الدور (طالب / عضو هيئة تدريس).",
            "بعد تسجيل الدخول الناجح يتم توجيه المستخدم تلقائيًا إلى لوحة المحادثة الذكية.",
            "آلية الأمان تعتمد على JWT (JSON Web Token) لإدارة الجلسات بشكل آمن.",
        ],
    ),
    (
        "01b_login_student_role.png",
        "1.1 تسجيل الدخول – اختيار دور الطالب",
        [
            "نفس الشاشة تتيح اختيار الدور (طالب) قبل الدخول؛ مما يحدد لاحقًا الصلاحيات والقوائم المعروضة.",
            "تجربة موحدة لكل الأدوار مع تخصيص ذكي لمحتوى الواجهة بعد الدخول.",
        ],
    ),
    (
        "02_register_faculty.png",
        "2. صفحة إنشاء حساب جديد (Register) – عضو هيئة تدريس",
        [
            "إنشاء حساب جديد لعضو هيئة تدريس مع التحقق من صحة البيانات (Email, Password Match, Strong Username).",
            "حقول النموذج: اسم المستخدم، البريد الإلكتروني، كلمة المرور وتأكيدها، اختيار الدور.",
            "كلمات المرور لا تُخزَّن نصيًا؛ يتم تشفيرها باستخدام bcrypt قبل الحفظ.",
        ],
    ),
    (
        "02b_register_student.png",
        "2.1 إنشاء حساب – طالب مع تخصيص",
        [
            "عند اختيار دور (طالب) تظهر قائمة التخصصات تلقائيًا (Computer Science, AI, Data Science, ...).",
            "التخصص يُستخدم لاحقًا لفلترة الكورسات الموصى بها وملفات المقررات الموجهة للطالب.",
        ],
    ),
    (
        "03_chat_faculty.png",
        "3. واجهة المحادثة الذكية (Chat) – عضو هيئة تدريس",
        [
            "الواجهة الأساسية للمساعد الذكي MANAMU بدعم كامل للغة العربية والإنجليزية.",
            "الشريط الجانبي يحتوي على: الصفحة الرئيسية، رفع الملفات، إدارة الكورسات، تصدير المحادثة PDF.",
            "زر Web Search يُفعِّل البحث الخارجي عند الحاجة لمعلومات خارج قاعدة المعرفة المحلية (RAG).",
            "عرض زمن الرسائل وإمكانية متابعة المحادثة بسلاسة (Auto-scroll).",
        ],
    ),
    (
        "03b_chat_faculty_upload_tab.png",
        "3.1 تبويب إدارة الملفات داخل المحادثة",
        [
            "أعضاء هيئة التدريس يمكنهم رفع ملفات PDF تخصّ المقرر داخل نفس الواجهة.",
            "يمكن تحديد الفئة المستهدفة (طلاب التخصص الفلاني) بحيث يستفيد منها الطلاب فقط.",
            "الملفات المرفوعة تتم فهرستها تلقائيًا داخل قاعدة بيانات شعاعية (PostgreSQL + pgvector) لاستخدامها في إجابات RAG.",
        ],
    ),
    (
        "04_faculty_courses.png",
        "4. إدارة الكورسات (Course Management) – عضو هيئة تدريس",
        [
            "لوحة احترافية لإدارة الكورسات: إنشاء، تعديل، حذف، وإظهار حالة الكورس (Active/Inactive).",
            "بطاقة كل كورس تعرض الصورة (Thumbnail) – يتم استخراجها تلقائيًا من YouTube عند توفر الرابط.",
            "إمكانية البحث والتصفية حسب التخصص + عداد للطلاب المسجلين بكل كورس.",
            "زر Visit Course لفتح الكورس الخارجي مباشرة في تبويب جديد.",
        ],
    ),
    (
        "05_chat_student.png",
        "5. واجهة المحادثة الذكية (Chat) – طالب",
        [
            "نفس واجهة المحادثة لكن بصلاحيات وقوائم مخصصة للطالب (لا تظهر له خيارات الرفع).",
            "الطالب يطرح أسئلته ويتلقى إجابات مدعومة بمحتوى المقررات المرفوعة من قِبل الدكتور.",
            "الإجابة تستخدم تقنية RAG (Retrieval-Augmented Generation) لرفع الدقة وتقليل الهلوسات.",
        ],
    ),
    (
        "06_student_courses_available.png",
        "6. منصة التعلم (Learning Platform) – الطالب",
        [
            "تعرض الكورسات الموصى بها بناءً على تخصص الطالب (في المثال: Computer Science).",
            "تبويبان: Available Courses (المتاحة) و My Enrollments (المسجَّل بها).",
            "زر Enroll Now لتسجيل الكورس مباشرة، وزر Preview لمعاينته قبل التسجيل.",
            "بحث وتصفية متقدمة لتسهيل الوصول للمحتوى المناسب.",
        ],
    ),
    (
        "07_upload_students.png",
        "7. رفع بيانات الطلاب من Excel (Admin)",
        [
            "أداة إدارية لإنشاء حسابات الطلاب دفعة واحدة من ملف Excel.",
            "متطلبات الملف موضّحة: الأعمدة المطلوبة (Name, University ID, Phone Number) – صيغة .xlsx أو .xls.",
            "Drag & Drop لرفع الملف بسهولة، مع تحقق فوري من صحة الصيغة.",
        ],
    ),
    (
        "08_upload_faculty.png",
        "8. رفع بيانات أعضاء هيئة التدريس من Excel (Admin)",
        [
            "نفس مفهوم الأداة السابقة لكن مخصصة لإنشاء حسابات أعضاء هيئة التدريس.",
            "يوفّر الكثير من الوقت عند بداية الفصل الدراسي مقارنة بالإدخال اليدوي.",
        ],
    ),
    (
        "09_upload_staff.png",
        "9. رفع بيانات الموظفين / المستخدمين الإداريين (Admin)",
        [
            "إنشاء حسابات للموظفين الإداريين (Superuser) من ملف Excel.",
            "صلاحيات أعلى لإدارة المستخدمين والمحتوى داخل النظام.",
        ],
    ),
]


# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------
def set_rtl_paragraph(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def set_run_rtl(run):
    """Mark a run as right-to-left so Word applies bidi correctly to mixed text."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)
    # Use a font that has both Arabic & Latin glyphs so cs/ascii align visually
    rfonts = rPr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rPr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:cs"), "Arial")
    # Ensure complex-script size matches the run size
    sz = rPr.find(qn("w:sz"))
    if sz is not None:
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), sz.get(qn("w:val")))
        rPr.append(szCs)


def set_doc_default_font(doc, font_name="Calibri", size_pt=12):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(size_pt)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def set_section_rtl(section):
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)


def add_centered_image(doc, path: Path, width_in=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))


def add_heading(doc, text, level=1, color=(31, 73, 125)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(p)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20 if level == 0 else 16 if level == 1 else 14)
    run.font.color.rgb = RGBColor(*color)
    set_run_rtl(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl_paragraph(p)
        run = p.add_run(item)
        run.font.size = Pt(12)
        set_run_rtl(run)


def add_paragraph_rtl(doc, text, bold=False, size=12, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_paragraph()
    p.alignment = align
    set_rtl_paragraph(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    set_run_rtl(run)
    return p


# --------------------------------------------------------------------------
# Build the document
# --------------------------------------------------------------------------
def build():
    doc = Document()
    set_doc_default_font(doc, "Calibri", 12)

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        set_section_rtl(section)

    # ----- Cover page ----------------------------------------------------
    for _ in range(3):
        doc.add_paragraph()

    add_paragraph_rtl(
        doc,
        "MANAMU",
        bold=True,
        size=48,
        color=(31, 73, 125),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph_rtl(
        doc,
        "المساعد التعليمي الذكي",
        bold=True,
        size=28,
        color=(0, 0, 0),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph_rtl(
        doc,
        "AI-Powered Learning Assistant",
        bold=False,
        size=16,
        color=(100, 100, 100),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for _ in range(2):
        doc.add_paragraph()
    add_paragraph_rtl(
        doc,
        "تقرير شرح كامل للنظام مع لقطات لكل الشاشات",
        bold=True,
        size=18,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph_rtl(
        doc,
        "إعداد: فريق MANAMU",
        bold=False,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_page_break()

    # ----- Overview ------------------------------------------------------
    add_heading(doc, "نظرة عامة على النظام", level=0)
    add_paragraph_rtl(
        doc,
        "MANAMU هو مساعد تعليمي ذكي يهدف إلى دعم الطلاب وأعضاء هيئة التدريس داخل البيئة "
        "الجامعية. يجمع النظام بين قوة نماذج اللغة الكبيرة (LLMs) وتقنية الاسترجاع المعزز "
        "بالتوليد (RAG) لتقديم إجابات دقيقة مبنية على محتوى المقررات الفعلي وليس على المعرفة "
        "العامة فقط، مع إضافة منصة كاملة لإدارة الكورسات وتسجيل الطلاب.",
    )
    add_heading(doc, "أبرز الميزات", level=1)
    add_bullets(
        doc,
        [
            "نظام مصادقة آمن باستخدام JWT وتشفير كلمات المرور بـ bcrypt.",
            "مساعد ذكي يعتمد RAG على ملفات المقرر المرفوعة (PDF) داخل قاعدة بيانات شعاعية (pgvector).",
            "إمكانية تفعيل البحث على الويب لتعزيز الإجابة عند نقص المعرفة الداخلية.",
            "منصة كورسات متكاملة: إنشاء/تعديل/حذف الكورسات، تسجيل الطلاب، استخراج صور YouTube تلقائيًا.",
            "أدوات إدارية لرفع بيانات الطلاب وأعضاء هيئة التدريس والموظفين دفعة واحدة من Excel.",
            "تصميم متجاوب وحديث يدعم العربية والإنجليزية مع تجربة مستخدم سلسة.",
            "تصدير المحادثات إلى PDF بسهولة.",
            "فصل واضح للأدوار: طالب – عضو هيئة تدريس – مسؤول، مع صلاحيات مخصّصة لكل دور.",
        ],
    )

    add_heading(doc, "البنية التقنية (Technology Stack)", level=1)
    add_bullets(
        doc,
        [
            "الواجهة الأمامية: React 18 + React Router + Axios + lucide-react + react-hook-form.",
            "الخادم الخلفي: FastAPI (Python) + SQLAlchemy + Pydantic.",
            "قاعدة البيانات: PostgreSQL مع امتداد pgvector لتخزين التضمينات الشعاعية.",
            "الذكاء الاصطناعي: OpenAI GPT + Sentence-Transformers + LangChain.",
            "البحث على الويب الاختياري عبر Tavily API.",
            "المصادقة: JWT + bcrypt.",
        ],
    )
    doc.add_page_break()

    # ----- Pages ---------------------------------------------------------
    add_heading(doc, "جولة كاملة في الشاشات", level=0)
    add_paragraph_rtl(
        doc,
        "الصفحات التالية تعرض كل شاشة من شاشات النظام مع شرح موجز ومركّز لوظيفتها وقيمتها للمستخدم.",
    )
    doc.add_paragraph()

    for fname, title, bullets in SECTIONS:
        img = SHOTS / fname
        if not img.exists():
            print(f"missing image: {fname}")
            continue
        add_heading(doc, title, level=1)
        add_bullets(doc, bullets)
        doc.add_paragraph()
        add_centered_image(doc, img, width_in=6.3)
        doc.add_page_break()

    # ----- Closing ------------------------------------------------------
    add_heading(doc, "الخلاصة", level=0)
    add_paragraph_rtl(
        doc,
        "يقدم نظام MANAMU حلًّا متكاملًا للتعليم الذكي يجمع بين مساعد محادثة قوي قائم على RAG، "
        "ومنصة لإدارة الكورسات، وأدوات إدارية تختصر الكثير من العمل اليدوي. تصميم النظام "
        "ومرونته يجعلانه جاهزًا للتوسع ليشمل كليات وأقسام أخرى داخل الجامعة بسهولة.",
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
